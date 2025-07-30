# app/routes.py

from flask import Blueprint, render_template, make_response, redirect, url_for, flash, request
from flask_login import login_user, logout_user, current_user, login_required
from flask import send_file
from docx import Document
from app.models import User
from app.forms import SignupForm, LoginForm
from app import db
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import login_required, current_user
from app.models import Resume
from app.forms import ResumeForm
from io import BytesIO
from reportlab.pdfgen import canvas
from weasyprint import HTML
import tempfile
import os
from werkzeug.utils import secure_filename
from PIL import Image
from transformers import pipeline
from huggingface_hub import login
from dotenv import load_dotenv
from huggingface_hub import HfApi
import re

load_dotenv()
login(token=os.getenv("HF_TOKEN"))
#api = HfApi()
#api.set_access_token(os.getenv("HF_TOKEN"))

routes = Blueprint('routes', __name__)

# Configure upload settings
UPLOAD_FOLDER = 'app/static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_profile_image(file, user_id):
    if file and allowed_file(file.filename):
        # Create upload directory if it doesn't exist
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        
        # Generate unique filename
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower()
        new_filename = f"profile_{user_id}_{current_user.id}.{file_ext}"
        filepath = os.path.join(UPLOAD_FOLDER, new_filename)
        
        # Save and resize image
        try:
            image = Image.open(file)
            # Resize to reasonable size (300x300 max)
            image.thumbnail((300, 300), Image.Resampling.LANCZOS)
            image.save(filepath, quality=85, optimize=True)
            return new_filename
        except Exception as e:
            print(f"Error processing image: {e}")
            return None
    return None

@routes.route('/signup', methods=['GET', 'POST'])
def signup():
    form = SignupForm()
    if form.validate_on_submit():
        # Create a new user instance
        user = User(
            username=form.username.data,
            email=form.email.data,
            password=generate_password_hash(form.password.data)
        )
        db.session.add(user)
        db.session.commit()

        flash('Signup successful! Please log in.', 'success')
        return redirect(url_for('routes.login'))  # 👈 redirect after signup

    return render_template('signup.html', form=form)

@routes.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and check_password_hash(user.password, form.password.data):
            login_user(user)
            flash('Logged in successfully!', 'success')
            return redirect(url_for('routes.resume'))
        else:
            flash('Invalid email or password.', 'danger')
    return render_template('login.html', form=form)

@routes.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out.', 'info')
    return redirect(url_for('routes.login'))

@routes.route('/resume', methods=['GET', 'POST'])
@login_required
def resume():
    saved_resume = Resume.query.filter_by(user_id=current_user.id).first()
    form = ResumeForm(obj=saved_resume)

    if form.validate_on_submit():
        if 'generate_ai' in request.form:
            # --- ✨ AI GENERATION ---
            prompt = f"""
            Create a professional resume for someone named {form.full_name.data or 'John Doe'}.
            Skills: {form.skills.data or 'Python, HTML, CSS'}
            Experience: {form.experience.data or '1 year web developer'}
            Education: {form.education.data or 'B.Tech in Computer Science'}
            """

            try:
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You're a helpful resume writer."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7
                )

                content = response['choices'][0]['message']['content']

                # Fill generated content into the experience field
                form.experience.data = content

                flash("AI-generated content added. Review and save your resume.", "info")
                return render_template("resume.html", form=form, resume=saved_resume)

            except OpenAIError as e:
                flash(f"AI generation failed: {str(e)}", "danger")
                return render_template("resume.html", form=form, resume=saved_resume)

        # ✅ Save logic
        if saved_resume:
            saved_resume.full_name = form.full_name.data
            saved_resume.email = form.email.data
            saved_resume.phone = form.phone.data
            saved_resume.address = form.address.data
            saved_resume.education = form.education.data
            saved_resume.skills = form.skills.data
            saved_resume.interpersonal_skills = form.interpersonal_skills.data
            saved_resume.experience = form.experience.data
            saved_resume.certificates = form.certificates.data
            saved_resume.github = form.github.data
            saved_resume.linkedin = form.linkedin.data
            saved_resume.layout = form.layout.data  # ✅ this line is critical

            # Save profile image if uploaded
            if form.profile_image.data:
                filename = save_profile_image(form.profile_image.data, current_user.id)
                if filename:
                    saved_resume.profile_image = filename

        else:
            resume = Resume(
                full_name=form.full_name.data,
                email=form.email.data,
                phone=form.phone.data,
                address=form.address.data,
                education=form.education.data,
                skills=form.skills.data,
                interpersonal_skills=form.interpersonal_skills.data,
                experience=form.experience.data,
                certificates=form.certificates.data,
                github=form.github.data,
                linkedin=form.linkedin.data,
                layout=form.layout.data,
                user_id=current_user.id
            )

            if form.profile_image.data:
                filename = save_profile_image(form.profile_image.data, current_user.id)
                if filename:
                    resume.profile_image = filename

            db.session.add(resume)

        db.session.commit()
        flash("Resume saved successfully.", "success")
        return redirect(url_for('routes.resume'))

    return render_template('resume.html', form=form, resume=saved_resume)

@routes.route('/download/pdf')
@login_required
def download_pdf():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        flash("No resume found.", "danger")
        return redirect(url_for('routes.resume'))

    buffer = BytesIO()
    p = canvas.Canvas(buffer)
    
    # Set up fonts and colors
    p.setFont("Helvetica-Bold", 24)
    y = 800
    
    # Title
    p.drawString(100, y, resume.full_name)
    y -= 40
    
    # Contact info
    p.setFont("Helvetica", 12)
    p.drawString(100, y, f"Email: {resume.email}")
    y -= 15
    p.drawString(100, y, f"Phone: {resume.phone}")
    y -= 15
    p.drawString(100, y, "Address:")
    y -= 15
    # Handle multi-line address
    for line in resume.address.split('\n'):
        if y < 50:  # Check if we need a new page
            p.showPage()
            y = 750
        p.drawString(120, y, line)
        y -= 15
    y -= 15
    
    # Social links
    if resume.github or resume.linkedin:
        p.setFont("Helvetica-Bold", 12)
        p.drawString(100, y, "Links:")
        y -= 15
        p.setFont("Helvetica", 12)
        if resume.github:
            p.drawString(120, y, f"GitHub: {resume.github}")
            y -= 15
        if resume.linkedin:
            p.drawString(120, y, f"LinkedIn: {resume.linkedin}")
            y -= 15
        y -= 15
    
    # Education section
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, y, "Education")
    y -= 20
    p.setFont("Helvetica", 12)
    # Handle multi-line text
    for line in resume.education.split('\n'):
        if y < 50:  # Check if we need a new page
            p.showPage()
            y = 750
        p.drawString(100, y, line)
        y -= 15
    y -= 15
    
    # Technical Skills section
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, y, "Technical Skills")
    y -= 20
    p.setFont("Helvetica", 12)
    # Handle multi-line text
    for line in resume.skills.split('\n'):
        if y < 50:  # Check if we need a new page
            p.showPage()
            y = 750
        p.drawString(100, y, line)
        y -= 15
    y -= 15
    
    # Interpersonal Skills section
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, y, "Interpersonal Skills")
    y -= 20
    p.setFont("Helvetica", 12)
    # Handle multi-line text
    for line in resume.interpersonal_skills.split('\n'):
        if y < 50:  # Check if we need a new page
            p.showPage()
            y = 750
        p.drawString(100, y, line)
        y -= 15
    y -= 15
    
    # Experience section
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, y, "Experience")
    y -= 20
    p.setFont("Helvetica", 12)
    # Handle multi-line text
    for line in resume.experience.split('\n'):
        if y < 50:  # Check if we need a new page
            p.showPage()
            y = 750
        p.drawString(100, y, line)
        y -= 15
    y -= 15
    
    # Certificates section (optional)
    if resume.certificates:
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, y, "Certificates")
        y -= 20
        p.setFont("Helvetica", 12)
        # Handle multi-line text
        for line in resume.certificates.split('\n'):
            if y < 50:  # Check if we need a new page
                p.showPage()
                y = 750
            p.drawString(100, y, line)
            y -= 15
    
    p.showPage()
    p.save()

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="resume.pdf", mimetype='application/pdf')


@routes.route('/download/pdf-html')
@login_required
def download_pdf_html():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        flash("No resume found.", "danger")
        return redirect(url_for('routes.resume'))

    # Choose template based on layout
    template_mapping = {
        'classic': 'resume_classic.html',
        'modern': 'resume_modern.html',
        'creative': 'resume_creative.html',
        'minimalist': 'resume_minimalist.html',
        'executive': 'resume_executive.html',
        'developer': 'resume_developer.html',
        'designer': 'resume_designer.html',
        'academic': 'resume_academic.html'
    }
    
    template = template_mapping.get(resume.layout, 'resume_classic.html')
    
    # Render the HTML template
    from flask import render_template_string
    html_content = render_template(template, resume=resume)
    
    # Convert HTML to PDF
    pdf = HTML(string=html_content).write_pdf()
    
    buffer = BytesIO()
    buffer.write(pdf)
    buffer.seek(0)
    
    return send_file(buffer, as_attachment=True, download_name="resume.pdf", mimetype='application/pdf')


@routes.route('/download/docx')
@login_required
def download_docx():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        flash("No resume found.", "danger")
        return redirect(url_for('routes.resume'))

    doc = Document()
    
    # Title
    title = doc.add_heading(resume.full_name, 0)
    title.alignment = 1  # Center alignment
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.add_run(f"Email: {resume.email}")
    contact_para.add_run('\n')
    contact_para.add_run(f"Phone: {resume.phone}")
    contact_para.add_run('\n')
    contact_para.add_run("Address:")
    contact_para.alignment = 1  # Center alignment
    
    # Address
    address_para = doc.add_paragraph()
    for line in resume.address.split('\n'):
        address_para.add_run(line)
        if line != resume.address.split('\n')[-1]:  # Not the last line
            address_para.add_run('\n')
    address_para.alignment = 1  # Center alignment
    
    # Social links
    if resume.github or resume.linkedin:
        links_para = doc.add_paragraph()
        if resume.github:
            links_para.add_run(f"GitHub: {resume.github}")
        if resume.github and resume.linkedin:
            links_para.add_run('\n')
        if resume.linkedin:
            links_para.add_run(f"LinkedIn: {resume.linkedin}")
        links_para.alignment = 1  # Center alignment
    
    # Add some spacing
    doc.add_paragraph()
    
    # Education section
    doc.add_heading('Education', level=1)
    education_para = doc.add_paragraph()
    for line in resume.education.split('\n'):
        education_para.add_run(line)
        if line != resume.education.split('\n')[-1]:  # Not the last line
            education_para.add_run('\n')
    
    # Technical Skills section
    doc.add_heading('Technical Skills', level=1)
    skills_para = doc.add_paragraph()
    for line in resume.skills.split('\n'):
        skills_para.add_run(line)
        if line != resume.skills.split('\n')[-1]:  # Not the last line
            skills_para.add_run('\n')
    
    # Interpersonal Skills section
    doc.add_heading('Interpersonal Skills', level=1)
    interpersonal_para = doc.add_paragraph()
    for line in resume.interpersonal_skills.split('\n'):
        interpersonal_para.add_run(line)
        if line != resume.interpersonal_skills.split('\n')[-1]:  # Not the last line
            interpersonal_para.add_run('\n')
    
    # Experience section
    doc.add_heading('Experience', level=1)
    experience_para = doc.add_paragraph()
    for line in resume.experience.split('\n'):
        experience_para.add_run(line)
        if line != resume.experience.split('\n')[-1]:  # Not the last line
            experience_para.add_run('\n')
    
    # Certificates section (optional)
    if resume.certificates:
        doc.add_heading('Certificates', level=1)
        certificates_para = doc.add_paragraph()
        for line in resume.certificates.split('\n'):
            certificates_para.add_run(line)
            if line != resume.certificates.split('\n')[-1]:  # Not the last line
                certificates_para.add_run('\n')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="resume.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@routes.route('/view_resume')
@login_required
def view_resume():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        flash('No resume found.', 'warning')
        return redirect(url_for('routes.resume'))

    template_mapping = {
        'classic': 'resume_classic.html',
        'modern': 'resume_modern.html',
        'creative': 'resume_creative.html',
        'minimalist': 'resume_minimalist.html',
        'executive': 'resume_executive.html',
        'developer': 'resume_developer.html',
        'designer': 'resume_designer.html',
        'academic': 'resume_academic.html'
    }
    
    template = template_mapping.get(resume.layout, 'resume_classic.html')
    return render_template(template, resume=resume)


#client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))  Set this in your environment
#HfApi().set_access_token(os.getenv("HF_TOKEN"))

generator = pipeline("text-generation",  model="EleutherAI/gpt-neo-1.3B", device=-1)

@routes.route('/generate_resume', methods=['GET', 'POST'])
@login_required
def generate_resume():
    if request.method == 'POST':
        job_title = request.form.get('job_title')
        skills = request.form.get('skills')
        experience = request.form.get('experience')

        # Clean prompt for summary
        summary_prompt = f"""
        You are a professional resume writer.
        Generate a 2-3 sentence summary for someone applying as a {job_title} with skills in {skills}.
        """

        experience_prompt = f"""
        Write 3 professional bullet points for resume experience based on the following role:
        {experience}
        """

        try:
            # Generate summary
            summary_result = generator(summary_prompt, max_new_tokens=100, temperature=0.7)[0]['generated_text']
            summary = summary_result.replace(summary_prompt, '').strip()

            # Generate experience
            exp_result = generator(experience_prompt, max_new_tokens=150, temperature=0.7)[0]['generated_text']
            exp_points = exp_result.replace(experience_prompt, '').strip()

            # OPTIONAL: clean garbage content from experience
            import re
            exp_points = re.sub(r"(Continue Reading|Share this|Subscribe|Tweet|Facebook).*", "", exp_points, flags=re.IGNORECASE).strip()

            return render_template(
                'generated_resume.html',
                summary=summary,
                experience=exp_points,
                job_title=job_title,
                skills=skills
            )

        except Exception as e:
            flash(f"Error: {e}", "danger")

    return render_template('generate_form.html')
